import os
import re
import logging
import zipfile
import xml.etree.ElementTree as ET
from typing import Dict, Any, Optional

logger = logging.getLogger("DocumentParserService")

class DocumentParserService:
    """
    Serviço universal de parsing e extração de texto de documentos para a Maeve.
    Suporta .pdf, .docx, .txt, .md, .csv, .json, .yaml e arquivos de código.
    Inclui detecção contextual inteligente para perfis profissionais / currículos.
    """

    CV_KEYWORDS = [
        "currículo", "curriculo", "resume", "curriculum vitae", "cv",
        "experiência profissional", "experiencia profissional", "experiências",
        "formação acadêmica", "formacao academica", "educação", "educacao",
        "habilidades técnicas", "stack tecnológica", "competências",
        "histórico profissional", "resumo profissional", "trajetória profissional"
    ]

    @classmethod
    def parse_document(cls, file_path: str, filename: str, mime_type: Optional[str] = None) -> Dict[str, Any]:
        """
        Extrai o conteúdo de texto de um arquivo no disco com base na extensão e/ou MIME type.
        Retorna dicionário com status, texto extraído, metadados e prompt contextual recomendado.
        """
        if not os.path.exists(file_path):
            return {"success": False, "error": f"Arquivo não encontrado: {file_path}"}

        lower_name = filename.lower()
        file_ext = os.path.splitext(lower_name)[1]

        try:
            # 1. Documentos PDF
            if file_ext == ".pdf" or (mime_type and "pdf" in mime_type):
                text = cls._extract_pdf(file_path)
                file_type = "pdf"

            # 2. Documentos DOCX (Word)
            elif file_ext in [".docx", ".doc"] or (mime_type and "wordprocessingml" in mime_type):
                text = cls._extract_docx(file_path)
                file_type = "docx"

            # 3. Arquivos de Texto / Markdown / Código
            elif file_ext in [".txt", ".md", ".json", ".csv", ".tsv", ".yaml", ".yml", ".xml", ".html", ".py", ".sql", ".log"] or (mime_type and mime_type.startswith("text/")):
                text = cls._extract_text(file_path)
                file_type = "text"

            else:
                # Tenta leitura textual como fallback resiliente
                try:
                    text = cls._extract_text(file_path)
                    file_type = "unknown_text"
                except Exception:
                    return {
                        "success": False,
                        "error": f"Formato não suportado para '{filename}'. Formatos aceitos: .docx, .pdf, .txt, .md, .csv, .json."
                    }

            clean_text = text.strip()
            if not clean_text:
                return {
                    "success": False,
                    "error": f"Não foi possível extrair nenhum texto legível de '{filename}' (o arquivo pode estar vazio ou conter apenas imagens escaneadas)."
                }

            # Identificação de Perfil Profissional / Currículo
            is_cv = cls._detect_cv(clean_text, lower_name)
            suggested_prompt = cls._build_prompt(filename, clean_text, is_cv)

            return {
                "success": True,
                "text": clean_text,
                "filename": filename,
                "file_type": file_type,
                "is_cv": is_cv,
                "char_count": len(clean_text),
                "suggested_prompt": suggested_prompt
            }

        except Exception as e:
            logger.error(f"Erro ao processar documento '{filename}': {e}", exc_info=True)
            return {"success": False, "error": f"Falha ao ler '{filename}': {str(e)}"}

    @classmethod
    def _extract_pdf(cls, file_path: str) -> str:
        """Extrai texto de PDFs usando pypdf."""
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        pages_text = []
        for i, page in enumerate(reader.pages):
            page_content = page.extract_text()
            if page_content:
                pages_text.append(page_content)
        return "\n\n".join(pages_text)

    @classmethod
    def _extract_docx(cls, file_path: str) -> str:
        """
        Extrai texto de arquivos .docx.
        Tenta primeiro python-docx; se não estiver instalado ou falhar,
        usa fallback nativo em stdlib (zipfile + xml parser).
        """
        # Tentativa 1: python-docx (preserva tabelas e parágrafos estruturados)
        try:
            import docx
            doc = docx.Document(file_path)
            content_parts = []
            
            # Parágrafos
            for p in doc.paragraphs:
                p_text = p.text.strip()
                if p_text:
                    content_parts.append(p_text)
            
            # Tabelas
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        content_parts.append(" | ".join(row_cells))
            
            extracted = "\n".join(content_parts)
            if extracted.strip():
                return extracted
        except Exception as e:
            logger.warning(f"python-docx falhou ou ausente, usando fallback XML nativo: {e}")

        # Tentativa 2: Fallback puro stdlib (zero dependências)
        try:
            with zipfile.ZipFile(file_path, "r") as docx_zip:
                xml_content = docx_zip.read("word/document.xml")
                tree = ET.fromstring(xml_content)
                namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                paragraphs = []
                for p in tree.iterfind(".//w:p", namespaces):
                    texts = [node.text for node in p.iterfind(".//w:t", namespaces) if node.text]
                    if texts:
                        paragraphs.append("".join(texts))
                return "\n".join(paragraphs)
        except Exception as fallback_err:
            raise RuntimeError(f"Falha na extração DOCX (docx e fallback): {fallback_err}")

    @classmethod
    def _extract_text(cls, file_path: str) -> str:
        """Lê arquivos de texto com detecção de encoding (utf-8, latin-1)."""
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
        for enc in encodings:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise ValueError(f"Não foi possível decodificar o arquivo de texto com os encodings testados: {encodings}")

    @classmethod
    def _detect_cv(cls, text: str, lower_filename: str) -> bool:
        """Determina se o documento é um currículo / perfil profissional."""
        for kw in ["curriculo", "currículo", "resume", "cv"]:
            if kw in lower_filename:
                return True

        lower_text = text[:3000].lower()
        matches = sum(1 for kw in cls.CV_KEYWORDS if kw in lower_text)
        return matches >= 2

    @classmethod
    def _build_prompt(cls, filename: str, text: str, is_cv: bool, max_preview_chars: int = 12000) -> str:
        """Gera prompt contextual apropriado para o cérebro da Maeve."""
        truncated_text = text[:max_preview_chars]
        if len(text) > max_preview_chars:
            truncated_text += f"\n\n[...conteúdo truncado para contexto; total de {len(text)} caracteres]"

        if is_cv:
            return (
                f"📄 **Recebi o Currículo / Perfil Profissional do Erik** no arquivo `{filename}`.\n\n"
                f"**Conteúdo Extraído:**\n\n{truncated_text}\n\n"
                "**Instruções para a Maeve:**\n"
                "1. Analise detalhadamente a trajetória profissional, formação acadêmica, stack técnica, projetos de destaque e conquistas do Erik.\n"
                "2. Crie ou atualize uma nota completa no Obsidian chamada 'Perfil Profissional - Erik' na pasta 'Recursos/Perfil/' (com tags, frontmatter e resumo executivo).\n"
                "3. Registre na sua memória de longo prazo (Qdrant e decisões) para utilizar esse contexto exato da carreira dele em todas as nossas conversas de trabalho, mestrado e projetos.\n"
                "4. Responda ao Erik com uma síntese afiada e calorosa do que você aprendeu com o currículo dele."
            )
        else:
            return (
                f"📑 **Recebi um documento:** `{filename}`.\n\n"
                f"**Conteúdo Extraído:**\n\n{truncated_text}\n\n"
                "**Instruções para a Maeve:**\n"
                "1. Sintetize os pontos centrais, decisões ou conceitos do documento.\n"
                "2. Se for um material relevante de estudo ou trabalho, crie uma nota estruturada no Obsidian na pasta adequada.\n"
                "3. Indexe os conhecimentos-chave na memória vetorial."
            )
